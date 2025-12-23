"""Documentation creation module for RPA framework."""

from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
import markdown

from ..core.logger import LoggerMixin


class DocsModule(LoggerMixin):
    """Handle document creation and templating."""

    def create_word(
        self,
        output_path: str,
        title: Optional[str] = None,
        content: Optional[str] = None,
    ) -> str:
        """Create a new Word document.

        Args:
            output_path: Output file path
            title: Optional document title
            content: Optional initial content

        Returns:
            Path to created document
        """
        doc = Document()

        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if content:
            doc.add_paragraph(content)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        self.logger.info(f"Created Word document: {output_path}")
        return output_path

    def create_from_template(
        self,
        template_path: str,
        output_path: str,
        data: Dict[str, Any],
    ) -> str:
        """Create a Word document from a template with variable substitution.

        Args:
            template_path: Path to template document
            output_path: Output file path
            data: Dictionary of variables to substitute

        Returns:
            Path to created document
        """
        doc = Document(template_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        self.logger.info(f"Created document from template: {output_path}")
        return output_path

    def add_table(
        self,
        doc_path: str,
        data: List[Dict[str, Any]],
        headers: Optional[List[str]] = None,
    ) -> str:
        """Add a table to an existing Word document.

        Args:
            doc_path: Path to existing document
            data: List of dicts for table rows
            headers: Optional custom headers (default: dict keys)

        Returns:
            Path to modified document
        """
        doc = Document(doc_path)

        if not data:
            return doc_path

        headers = headers or list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            cell.paragraphs[0].runs[0].bold = True

        # Add data rows
        for row_data in data:
            row = table.add_row()
            for i, header in enumerate(headers):
                row.cells[i].text = str(row_data.get(header, ""))

        doc.save(doc_path)
        self.logger.info(f"Added table with {len(data)} rows to {doc_path}")
        return doc_path

    def add_image(
        self,
        doc_path: str,
        image_path: str,
        width: Optional[float] = None,
        caption: Optional[str] = None,
    ) -> str:
        """Add an image to a Word document.

        Args:
            doc_path: Path to document
            image_path: Path to image file
            width: Image width in inches
            caption: Optional caption text

        Returns:
            Path to modified document
        """
        doc = Document(doc_path)

        if width:
            doc.add_picture(image_path, width=Inches(width))
        else:
            doc.add_picture(image_path)

        if caption:
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(doc_path)
        self.logger.info(f"Added image to {doc_path}")
        return doc_path

    def create_markdown(
        self,
        output_path: str,
        title: str,
        sections: List[Dict[str, str]],
    ) -> str:
        """Create a Markdown document.

        Args:
            output_path: Output file path
            title: Document title
            sections: List of {'heading': str, 'content': str} dicts

        Returns:
            Path to created document
        """
        lines = [f"# {title}", ""]

        for section in sections:
            heading = section.get("heading", "")
            content = section.get("content", "")
            level = section.get("level", 2)

            if heading:
                lines.append(f"{'#' * level} {heading}")
                lines.append("")

            if content:
                lines.append(content)
                lines.append("")

        md_content = "\n".join(lines)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_text(md_content)
        self.logger.info(f"Created Markdown document: {output_path}")
        return output_path

    def markdown_to_html(
        self,
        input_path: str,
        output_path: Optional[str] = None,
        title: Optional[str] = None,
    ) -> str:
        """Convert Markdown to HTML.

        Args:
            input_path: Path to Markdown file
            output_path: Output HTML path (default: same name with .html)
            title: HTML page title

        Returns:
            Path to HTML file
        """
        md_content = Path(input_path).read_text()
        html_body = markdown.markdown(md_content, extensions=["tables", "fenced_code"])

        title = title or Path(input_path).stem

        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
               line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background: #f4f4f4; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

        if not output_path:
            output_path = str(Path(input_path).with_suffix(".html"))

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_text(html_content)
        self.logger.info(f"Converted Markdown to HTML: {output_path}")
        return output_path

    def create_html_report(
        self,
        output_path: str,
        title: str,
        data: Dict[str, Any],
        template: Optional[str] = None,
    ) -> str:
        """Create an HTML report from data.

        Args:
            output_path: Output file path
            title: Report title
            data: Data to include in report
            template: Optional custom Jinja2 template string

        Returns:
            Path to created report
        """
        if template:
            tmpl = Template(template)
            html_content = tmpl.render(title=title, data=data, now=datetime.now())
        else:
            # Default report template
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
        .timestamp {{ color: #666; font-size: 0.9em; }}
        .section {{ margin: 20px 0; padding: 15px; background: #f9f9f9; border-radius: 5px; }}
        .key {{ font-weight: bold; color: #555; }}
        .value {{ margin-left: 10px; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="timestamp">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    <div class="section">
"""
            for key, value in data.items():
                if isinstance(value, (list, dict)):
                    value = str(value)
                html_content += f'        <p><span class="key">{key}:</span> <span class="value">{value}</span></p>\n'

            html_content += """    </div>
</body>
</html>"""

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_text(html_content)
        self.logger.info(f"Created HTML report: {output_path}")
        return output_path

    def render_template(
        self,
        template_string: str,
        data: Dict[str, Any],
    ) -> str:
        """Render a Jinja2 template string with data.

        Args:
            template_string: Jinja2 template string
            data: Variables for template

        Returns:
            Rendered string
        """
        tmpl = Template(template_string)
        return tmpl.render(**data)

    def render_template_file(
        self,
        template_path: str,
        output_path: str,
        data: Dict[str, Any],
    ) -> str:
        """Render a Jinja2 template file with data.

        Args:
            template_path: Path to template file
            output_path: Output file path
            data: Variables for template

        Returns:
            Path to rendered file
        """
        template_content = Path(template_path).read_text()
        rendered = self.render_template(template_content, data)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_text(rendered)
        self.logger.info(f"Rendered template to: {output_path}")
        return output_path

    def extract_text_from_word(self, doc_path: str) -> str:
        """Extract all text from a Word document.

        Args:
            doc_path: Path to Word document

        Returns:
            Extracted text
        """
        doc = Document(doc_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    def fill_form(
        self,
        doc_path: str,
        output_path: str,
        field_mappings: List[Dict[str, Any]],
    ) -> str:
        """Fill form fields in a Word document by table cell positions.

        Args:
            doc_path: Path to source Word document
            output_path: Path to save filled document
            field_mappings: List of dicts with keys:
                - table: Table index (0-based)
                - row: Row index (0-based)
                - col: Column index (0-based)
                - value: Value to insert

        Returns:
            Path to filled document
        """
        doc = Document(doc_path)

        # Remove any document protection
        self._remove_protection(doc)

        for mapping in field_mappings:
            table_idx = mapping.get("table", 0)
            row_idx = mapping.get("row")
            col_idx = mapping.get("col")
            value = mapping.get("value", "")

            if row_idx is None or col_idx is None:
                continue

            try:
                table = doc.tables[table_idx]
                cell = table.rows[row_idx].cells[col_idx]

                # Clear existing content and add new editable text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = ""

                # Add value as a new run to ensure it's editable
                if cell.paragraphs:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = str(value)
                    else:
                        cell.paragraphs[0].add_run(str(value))
                else:
                    cell.text = str(value)

                self.logger.debug(f"Filled T{table_idx}[{row_idx},{col_idx}] = {value}")
            except (IndexError, AttributeError) as e:
                self.logger.warning(f"Could not fill T{table_idx}[{row_idx},{col_idx}]: {e}")

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        self.logger.info(f"Filled form and saved to: {output_path}")
        return output_path

    def _remove_protection(self, doc: Document) -> None:
        """Remove document protection to make it fully editable."""
        from docx.oxml.ns import qn

        # Access the document's XML
        doc_elm = doc.part.element

        # Find and remove documentProtection from settings
        for settings in doc_elm.iter(qn('w:settings')):
            for protection in settings.findall(qn('w:documentProtection')):
                settings.remove(protection)

        # Also check document root for protection
        for protection in doc_elm.iter(qn('w:documentProtection')):
            protection.getparent().remove(protection)

    def get_form_structure(self, doc_path: str) -> List[Dict[str, Any]]:
        """Get the structure of tables in a Word document for form filling.

        Args:
            doc_path: Path to Word document

        Returns:
            List of table structures with cell positions and content
        """
        doc = Document(doc_path)
        structure = []

        for t_idx, table in enumerate(doc.tables):
            table_info = {
                "table_index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "cells": []
            }

            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_info = {
                        "row": r_idx,
                        "col": c_idx,
                        "text": cell.text.strip()[:100],
                        "is_empty": not cell.text.strip()
                    }
                    table_info["cells"].append(cell_info)

            structure.append(table_info)

        return structure
